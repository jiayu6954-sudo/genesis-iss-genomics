"""
Convert manuscript_draft_v1.md → submission-ready .docx for bioRxiv.
Uses python-docx with proper academic formatting.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path('e:/miniconda3/envs/llama-env/genesis_project')
MD_FILE  = PROJECT_ROOT / 'science_engine/reports/manuscript_draft_v1.md'
FIG_DIR  = PROJECT_ROOT / 'science_engine/figures'
OUT_DOCX = PROJECT_ROOT / 'science_engine/reports/manuscript_draft_v1.docx'


def normalise(text: str) -> str:
    sup = {'⁰':'0','¹':'1','²':'2','³':'3','⁴':'4',
           '⁵':'5','⁶':'6','⁷':'7','⁸':'8','⁹':'9','⁻':'-',
           '₀':'0','₁':'1','₂':'2','₃':'3','₄':'4',
           '₅':'5','₆':'6','₇':'7','₈':'8','₉':'9'}
    for k, v in sup.items():
        text = text.replace(k, v)
    text = text.replace('\u2212', '-')
    return text


def add_line_numbers(doc):
    """Enable continuous line numbering in the document."""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:countBy'), '5')
        lnNumType.set(qn('w:restart'), 'continuous')
        sectPr.append(lnNumType)


def set_page_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)


def add_run_styled(para, text: str, bold=False, italic=False,
                   size=None, color=None):
    """Add a run with mixed bold/italic to a paragraph."""
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def render_inline(para, text: str, base_size=11):
    """Parse **bold**, *italic*, ***bold-italic*** and add runs."""
    text = normalise(text)
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*'   # bold-italic
        r'|\*\*(.+?)\*\*'        # bold
        r'|\*([^*\n]+?)\*)'      # italic
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            add_run_styled(para, text[last:m.start()], size=base_size)
        if m.group(2):
            add_run_styled(para, m.group(2), bold=True, italic=True, size=base_size)
        elif m.group(3):
            add_run_styled(para, m.group(3), bold=True, size=base_size)
        elif m.group(4):
            add_run_styled(para, m.group(4), italic=True, size=base_size)
        last = m.end()
    if last < len(text):
        add_run_styled(para, text[last:], size=base_size)


def render_table(doc, rows):
    """Render a markdown table as a Word table."""
    if len(rows) < 3:
        return
    header = rows[0]
    body   = [r for r in rows[2:] if any(c.strip() for c in r)]
    n_cols = len(header)
    if n_cols == 0:
        return

    table = doc.add_table(rows=1 + len(body), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(normalise(re.sub(r'\*+([^*]+)\*+', r'\1', cell_text.strip())))
        run.bold = True
        run.font.size = Pt(9)
        # Blue header fill
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'DCE6FF')
        tcPr.append(shd)

    # Body rows
    for ri, row in enumerate(body):
        row = (row + [''] * n_cols)[:n_cols]
        cells = table.rows[ri + 1].cells
        for i, cell_text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            render_inline(p, cell_text.strip(), base_size=9)

    doc.add_paragraph()  # spacing after table


def embed_figure(doc, png_path: Path, label: str, caption: str):
    """Add a figure page with embedded PNG and caption."""
    doc.add_page_break()
    # Label
    lbl_p = doc.add_paragraph()
    lbl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lbl_p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(20, 60, 120)

    # Image
    if png_path.exists():
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_p.add_run()
        img_run.add_picture(str(png_path), width=Inches(5.8))
    else:
        doc.add_paragraph(f'[Figure file not found: {png_path.name}]')

    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    render_inline(cap_p, caption, base_size=9)
    cap_p.paragraph_format.space_before = Pt(6)


def build_docx(md_path: Path, out_path: Path):
    text  = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    doc = Document()
    set_page_margins(doc)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    i            = 0
    in_abstract  = False
    abs_buf      = []
    tbl_buf      = []
    fig_done     = False
    author_done  = False

    fig_data = [
        (FIG_DIR / 'Figure1.png',  'Figure 1',
         'Figure 1. Species-stratified analysis of CRISPR-Cas and IS element '
         'distributions in ISS versus ground Bacillales. '
         '(a) CRISPR-Cas prevalence by species and environment; only '
         'P. polymyxa shows a significant ISS-vs-ground difference '
         '(Fisher p = 0.0095; BH q = 0.025). '
         '(b) Boxplots of IS element density (IS per 1,000 CDS) by species '
         'and environment. ISS = orange, Ground = blue.'),
        (FIG_DIR / 'Figure2.png',  'Figure 2',
         'Figure 2. Co-loss of CRISPR-Cas and IS elements in '
         'Paenibacillus polymyxa and metagenomics validation. '
         '(a) IS element density vs CRISPR-Cas status '
         '(p = 0.003, Cliff delta = -0.75). '
         '(b) CRISPR-associated KO density in GLDS-224 ISS debris vs '
         'ground debris (ratio 0.503-fold).'),
        (FIG_DIR / 'FigureS1.png', 'Supplementary Figure S1',
         'Supplementary Figure S1. Assembly quality assessment. '
         '(a) N50 vs IS density across all 85 genomes (r = +0.428). '
         '(b) IS per 1,000 CDS vs IS per Mbp normalisation concordance.'),
    ]

    while i < len(lines):
        raw = lines[i].rstrip()

        # Flush table buffer
        if not raw.startswith('|') and tbl_buf:
            render_table(doc, tbl_buf)
            tbl_buf = []

        # Horizontal rule
        if raw == '---':
            if in_abstract and abs_buf:
                full = normalise(' '.join(abs_buf))
                abs_p = doc.add_paragraph()
                abs_p.style = doc.styles['Normal']
                # Light blue shading
                pPr = abs_p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6FF')
                pPr.append(shd)
                abs_p.paragraph_format.left_indent  = Cm(0.5)
                abs_p.paragraph_format.right_indent = Cm(0.5)
                render_inline(abs_p, ' '.join(abs_buf), base_size=10)
                abs_p.paragraph_format.space_after = Pt(8)
                in_abstract = False
                abs_buf = []
            else:
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Insert figures before Figure Legends
        if not fig_done and raw.startswith('## Figure Legends'):
            for png, lbl, cap in fig_data:
                embed_figure(doc, png, lbl, cap)
            fig_done = True
            doc.add_page_break()

        # Table rows
        if raw.startswith('|'):
            tbl_buf.append([c.strip() for c in raw.strip('|').split('|')])
            i += 1
            continue

        # Headings
        if raw.startswith('#### '):
            p = doc.add_heading(normalise(raw[5:]), level=4)
            p.runs[0].font.size = Pt(10)
            i += 1; continue
        if raw.startswith('### '):
            p = doc.add_heading(normalise(raw[4:]), level=3)
            i += 1; continue
        if raw.startswith('## '):
            title = raw[3:]
            p = doc.add_heading(normalise(title), level=2)
            if title.strip() == 'Abstract':
                in_abstract = True
            i += 1; continue
        if raw.startswith('# ') and not raw.startswith('## '):
            p = doc.add_heading(normalise(raw[2:]), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1; continue

        # Author / Correspondence
        if raw.startswith('**Author:**'):
            if not author_done:
                ap = doc.add_paragraph()
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r1 = ap.add_run('ZJY')
                r1.bold = True; r1.font.size = Pt(12)
                ap.add_run('\nChina').italic = True
                ap.runs[-1].font.size = Pt(10)
                author_done = True
            i += 1; continue
        if raw.startswith('**Correspondence:**'):
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.add_run('Correspondence: jiayu6954@gmail.com').font.size = Pt(10)
            i += 1; continue

        # Running title / Keywords
        if any(raw.startswith(f'**{kw}') for kw in ['Running title', 'Keywords']):
            kp = doc.add_paragraph()
            render_inline(kp, raw, base_size=9)
            i += 1; continue

        # Numbered reference
        if re.match(r'^\d+\.', raw):
            if in_abstract:
                abs_buf.append(raw)
            else:
                rp = doc.add_paragraph(style='Normal')
                rp.paragraph_format.left_indent  = Cm(0.6)
                rp.paragraph_format.first_line_indent = Cm(-0.6)
                rp.paragraph_format.space_after   = Pt(2)
                render_inline(rp, raw, base_size=9)
            i += 1; continue

        # Bullet
        if raw.startswith('- ') or raw.startswith('* '):
            bp = doc.add_paragraph(style='List Bullet')
            render_inline(bp, raw[2:], base_size=10)
            i += 1; continue

        # Empty line
        if raw.strip() == '':
            if not in_abstract:
                pass  # paragraph spacing handles gaps
            i += 1; continue

        # Abstract accumulator
        if in_abstract:
            abs_buf.append(raw)
            i += 1; continue

        # Regular paragraph — collect continuation lines
        para_lines = [raw]
        while i + 1 < len(lines):
            nxt = lines[i + 1].rstrip()
            if (not nxt
                    or nxt.startswith('#')
                    or nxt.startswith('|')
                    or nxt == '---'
                    or nxt.startswith('- ')
                    or nxt.startswith('* ')
                    or re.match(r'^\d+\.\s+[A-Z]', nxt)
                    or nxt.startswith('**Running')
                    or nxt.startswith('**Keywords')
                    or nxt.startswith('**Author')
                    or nxt.startswith('**Correspondence')):
                break
            para_lines.append(nxt)
            i += 1

        full_para = ' '.join(para_lines)
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_after = Pt(6)
        render_inline(p, full_para, base_size=11)
        i += 1

    # Flush remaining table
    if tbl_buf:
        render_table(doc, tbl_buf)

    add_line_numbers(doc)
    doc.save(str(out_path))
    size_kb = round(out_path.stat().st_size / 1024)
    print(f'DOCX saved : {out_path}')
    print(f'Size       : {size_kb} KB')


if __name__ == '__main__':
    build_docx(MD_FILE, OUT_DOCX)
